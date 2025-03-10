
import docx
import docx.styles
from docx.shared import Pt

# Создаем новый документ
doc = docx.Document()

# Задаем стиль шрифта для документа
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(14)

# Добавляем абзац с маркировкой
doc.add_paragraph('Абзац!', 'List Bullet')

# Добавляем абзацы
par1 = doc.add_paragraph('first ')
par2 = doc.add_paragraph('first 22222 ')
par3 = doc.add_paragraph('first 3333 ')

# Добавляем форматированные текстовые фрагменты
par1.add_run('New latter in ab 1111').italic = True
par2.add_run('New latter in ab 2222222').bold = True
run = par3.add_run('News in 3-nd paragraph')
run.bold = True
run.underline = True

#Сохраняем
doc.save('example.docx')
